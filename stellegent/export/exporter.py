"""Export OCR + NLP results to DOCX, PDF, TXT, PNG, JSON manifest."""
from __future__ import annotations
import json
import re
from dataclasses import dataclass, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Sequence
import uuid

import cv2
import numpy as np

from ..ocr.engine import OCRLine


@dataclass
class ExportResult:
    lecture_id: str
    captured_at: str
    dir: str
    image_path: str
    txt_path: str
    docx_path: str
    pdf_path: str
    manifest_path: str

    def to_dict(self) -> dict:
        return asdict(self)


def _detect_tags(text: str) -> List[str]:
    tags: List[str] = []
    if re.search(r"[=∑∫√≈≠≤≥±π]|\b\d+\s*[+\-*/^]\s*\d+", text):
        tags.append("math-equations")
    if re.search(r"\b(def|class|return|import|function|const|let|var)\b", text):
        tags.append("code")
    if re.search(r"\b(theorem|proof|lemma|corollary)\b", text, re.I):
        tags.append("proofs")
    if re.search(r"^\s*\d+[.\)]\s", text, re.M):
        tags.append("numbered-steps")
    if re.search(r"^\s*[-•*]\s", text, re.M):
        tags.append("bullets")
    return tags


def _write_txt(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def _write_docx(path: Path, lines: Sequence[OCRLine], corrected: str,
                summary: str) -> None:
    from docx import Document  # type: ignore
    doc = Document()
    doc.add_heading("Lecture Capture", level=1)
    doc.add_heading("Summary", level=2)
    for ln in summary.splitlines():
        s = ln.strip()
        if not s:
            continue
        if s.startswith(("-", "•", "*")):
            doc.add_paragraph(s.lstrip("-•* ").strip(), style="List Bullet")
        else:
            doc.add_paragraph(s)
    doc.add_heading("Corrected Transcript", level=2)
    for para in corrected.split("\n"):
        if not para.strip():
            continue
        if re.match(r"^\s*\d+[.\)]\s", para):
            doc.add_paragraph(para, style="List Number")
        elif para.strip().startswith(("-", "•", "*")):
            doc.add_paragraph(para.lstrip("-•* "), style="List Bullet")
        else:
            doc.add_paragraph(para)
    doc.add_heading("Raw OCR (low-confidence flagged)", level=2)
    for ln in lines:
        marker = "  [LOW]" if ln.confidence < 0.75 else ""
        doc.add_paragraph(f"{ln.text}{marker} ({ln.confidence:.2f})")
    doc.save(str(path))


def _write_pdf(path: Path, image_path: Path, summary: str, corrected: str) -> None:
    from fpdf import FPDF  # type: ignore
    pdf = FPDF(unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    margin = 10
    page_w = pdf.w - 2 * margin
    pdf.set_left_margin(margin)
    pdf.set_right_margin(margin)

    def safe(s: str) -> str:
        return s.encode("latin-1", "replace").decode("latin-1")

    def write_block(lines: list, line_h: float) -> None:
        for ln in lines:
            text = safe(ln).strip()
            if not text:
                pdf.ln(line_h)
                continue
            pdf.multi_cell(page_w, line_h, text)

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(page_w, 10, safe("Lecture Capture"), ln=True)
    if image_path.exists():
        try:
            pdf.image(str(image_path), x=margin, y=25, w=page_w)
            pdf.ln(120)
        except Exception:
            pass
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(page_w, 8, "Summary", ln=True)
    pdf.set_font("Helvetica", "", 11)
    write_block(summary.splitlines(), 6)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(page_w, 8, "Corrected Transcript", ln=True)
    pdf.set_font("Helvetica", "", 10)
    write_block(corrected.splitlines(), 5)
    pdf.output(str(path))


def export_all(image: np.ndarray, lines: Sequence[OCRLine],
               corrected: str, summary: str, base_dir: Path,
               course_name: Optional[str] = None,
               lecture_id: Optional[str] = None) -> ExportResult:
    now = datetime.now(timezone.utc)
    lecture_id = lecture_id or str(uuid.uuid4())
    date_str = now.date().isoformat()
    out_dir = Path(base_dir) / date_str / lecture_id
    out_dir.mkdir(parents=True, exist_ok=True)

    image_path = out_dir / "board.png"
    cv2.imwrite(str(image_path), image)

    raw_text = "\n".join(l.text for l in lines)
    txt_path = out_dir / "transcript.txt"
    _write_txt(txt_path, raw_text)

    corrected_path = out_dir / "corrected.txt"
    _write_txt(corrected_path, corrected)

    summary_path = out_dir / "summary.txt"
    _write_txt(summary_path, summary)

    docx_path = out_dir / "lecture.docx"
    _write_docx(docx_path, lines, corrected, summary)

    pdf_path = out_dir / "lecture.pdf"
    _write_pdf(pdf_path, image_path, summary, corrected)

    manifest: Dict = {
        "lecture_id": lecture_id,
        "course_name": course_name,
        "captured_at": now.isoformat(),
        "language": "en",
        "engine": "stellegent",
        "tags": _detect_tags(raw_text + "\n" + corrected),
        "lines": [l.to_dict() for l in lines],
        "files": {
            "image": image_path.name,
            "txt": txt_path.name,
            "corrected": corrected_path.name,
            "summary": summary_path.name,
            "docx": docx_path.name,
            "pdf": pdf_path.name,
        },
    }
    manifest_path = out_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    return ExportResult(
        lecture_id=lecture_id,
        captured_at=now.isoformat(),
        dir=str(out_dir),
        image_path=str(image_path),
        txt_path=str(txt_path),
        docx_path=str(docx_path),
        pdf_path=str(pdf_path),
        manifest_path=str(manifest_path),
    )
