"""Export OCR + NLP results to DOCX, PDF, TXT, WebP images, JSON manifest.

DOCX, PDF and TXT share one layout: summary bullet points on top, the
transcript below. No images are embedded in any document. Board images are
saved separately as compressed WebP (raw + preprocessed) for the UI toggle.

The PDF is rendered in pure Python (fpdf2) from the SAME layout as the DOCX, so
the two stay 1:1 without any external program (no LibreOffice/Word). A bundled
DejaVu TrueType font is embedded so unicode (bullets, circled digits, math
symbols) renders identically on every platform, including Docker.
"""
from __future__ import annotations
import json
import re
from dataclasses import dataclass, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Sequence
import uuid

import cv2
import numpy as np

from ..ocr.base import OCRLine

# WebP encode quality (0-100). 80 keeps board text crisp at a fraction of PNG size.
_WEBP_QUALITY = 80

# Bundled unicode font (embedded into the PDF). Regular + bold mirror the DOCX
# normal/heading runs.
_FONTS_DIR = Path(__file__).parent / "assets" / "fonts"
_FONT_REGULAR = _FONTS_DIR / "DejaVuSans.ttf"
_FONT_BOLD = _FONTS_DIR / "DejaVuSans-Bold.ttf"

# Page geometry (mm). A4 with 1-inch (25.4 mm) margins mirrors Word defaults so
# the PDF and DOCX line up. Heading/body point sizes match the DOCX styles.
_PAGE_MARGIN = 25.4
_HEADING_PT = 16
_BODY_PT = 11
_LINE_H = 6.0  # mm per text line at body size


@dataclass
class ExportResult:
    lecture_id: str
    captured_at: str
    dir: str
    image_path: str
    raw_image_path: str
    txt_path: str
    docx_path: str
    pdf_path: str
    manifest_path: str

    def to_dict(self) -> dict:
        return asdict(self)


def _detect_tags(text: str) -> List[str]:
    tags: List[str] = []
    if re.search(r"[=∑∫√≈≠≤≥±π]|\b\d+\s*[+\-*/^]\s*\d+", text):
        tags.append("math-equations")
    if re.search(r"\b(def|class|return|import|function|const|let|var)\b", text):
        tags.append("code")
    if re.search(r"\b(theorem|proof|lemma|corollary)\b", text, re.I):
        tags.append("proofs")
    if re.search(r"^\s*\d+[.\)]\s", text, re.M):
        tags.append("numbered-steps")
    if re.search(r"^\s*[-•*]\s", text, re.M):
        tags.append("bullets")
    return tags


def _save_webp(path: Path, image: np.ndarray) -> None:
    cv2.imwrite(str(path), image, [cv2.IMWRITE_WEBP_QUALITY, _WEBP_QUALITY])


def _summary_bullets(summary: str) -> List[str]:
    """One clean bullet per non-empty line, list markers stripped."""
    out: List[str] = []
    for ln in summary.splitlines():
        s = ln.strip().lstrip("-•*").strip()
        if s:
            out.append(s)
    return out


def _transcript_paras(corrected: str) -> List[str]:
    return [p for p in corrected.split("\n") if p.strip()]


def _write_txt(path: Path, text: str) -> None:
    # utf-8-sig: prepend a BOM so Windows apps (Notepad et al.) detect the
    # encoding and open the file instead of choking on the unicode bytes
    # (bullets, circled digits, math symbols). Also normalise to CRLF so the
    # downloaded .txt opens cleanly in Windows editors.
    path.write_text(text.replace("\n", "\r\n"), encoding="utf-8-sig", newline="")


def _compose_document(summary: str, corrected: str) -> str:
    """Plain-text rendering shared shape: summary bullets, then transcript."""
    parts: List[str] = ["Summary", ""]
    bullets = _summary_bullets(summary)
    parts.extend(f"- {b}" for b in bullets) if bullets else parts.append("- (none)")
    parts += ["", "Transcript", "", corrected.strip()]
    return "\n".join(parts).rstrip() + "\n"


def _write_docx(path: Path, summary: str, corrected: str) -> None:
    from docx import Document  # type: ignore
    doc = Document()
    doc.add_heading("Summary", level=1)
    for b in _summary_bullets(summary):
        doc.add_paragraph(f"- {b}")
    doc.add_heading("Transcript", level=1)
    for para in _transcript_paras(corrected):
        doc.add_paragraph(para)
    doc.save(str(path))


def _write_pdf(path: Path, summary: str, corrected: str) -> None:
    """Render the PDF in pure Python, mirroring the DOCX layout exactly.

    Same content shape as `_write_docx`: a "Summary" heading + "- " bullets,
    then a "Transcript" heading + paragraphs. The bundled DejaVu font is
    embedded so all unicode renders faithfully (no latin-1 '?' degradation), and
    fpdf2 wraps every line to the page width + auto page-breaks, so nothing
    bleeds off the paper.
    """
    from fpdf import FPDF, XPos, YPos  # type: ignore

    pdf = FPDF(unit="mm", format="A4")
    pdf.set_margins(_PAGE_MARGIN, _PAGE_MARGIN, _PAGE_MARGIN)
    pdf.set_auto_page_break(auto=True, margin=_PAGE_MARGIN)
    pdf.add_font("DejaVu", "", str(_FONT_REGULAR))
    pdf.add_font("DejaVu", "B", str(_FONT_BOLD))
    pdf.add_page()
    epw = pdf.epw  # effective page width (inside margins)

    def heading(text: str) -> None:
        pdf.set_font("DejaVu", "B", _HEADING_PT)
        pdf.multi_cell(epw, _LINE_H * 1.4, text,
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    def body(text: str) -> None:
        pdf.set_font("DejaVu", "", _BODY_PT)
        pdf.multi_cell(epw, _LINE_H, text,
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    heading("Summary")
    bullets = _summary_bullets(summary)
    if bullets:
        for b in bullets:
            body(f"- {b}")
    else:
        body("- (none)")
    pdf.ln(3)
    heading("Transcript")
    for para in _transcript_paras(corrected):
        body(para)

    pdf.output(str(path))


def write_documents(*, docx_path: str | Path, pdf_path: str | Path,
                    txt_path: str | Path, summary: str, corrected: str) -> None:
    """(Re)write the three downloadable documents from text only.

    All three share one layout. The PDF is rendered in pure Python from the same
    content as the DOCX (see `_write_pdf`), so they stay 1:1 with no external
    program involved.
    """
    docx_path, pdf_path, txt_path = Path(docx_path), Path(pdf_path), Path(txt_path)
    _write_txt(txt_path, _compose_document(summary, corrected))
    _write_docx(docx_path, summary, corrected)
    _write_pdf(pdf_path, summary, corrected)


def export_all(image: np.ndarray, lines: Sequence[OCRLine],
               corrected: str, summary: str, base_dir: Path,
               course_name: Optional[str] = None,
               lecture_id: Optional[str] = None,
               raw_text: Optional[str] = None,
               raw_image: Optional[np.ndarray] = None) -> ExportResult:
    now = datetime.now(timezone.utc)
    lecture_id = lecture_id or str(uuid.uuid4())
    date_str = now.date().isoformat()
    out_dir = Path(base_dir) / date_str / lecture_id
    out_dir.mkdir(parents=True, exist_ok=True)

    # Preprocessed (default view) + raw (original capture), both compressed WebP.
    image_path = out_dir / "board.webp"
    _save_webp(image_path, image)
    raw_image_path = out_dir / "board_raw.webp"
    _save_webp(raw_image_path, image if raw_image is None else raw_image)

    # text engines (Gemini) carry no per-line list, so accept raw_text directly
    if raw_text is None:
        raw_text = "\n".join(l.text for l in lines)

    txt_path = out_dir / "transcript.txt"
    docx_path = out_dir / "lecture.docx"
    pdf_path = out_dir / "lecture.pdf"
    write_documents(docx_path=docx_path, pdf_path=pdf_path, txt_path=txt_path,
                    summary=summary, corrected=corrected)

    manifest: Dict = {
        "lecture_id": lecture_id,
        "course_name": course_name,
        "captured_at": now.isoformat(),
        "language": "en",
        "engine": "stellegent",
        "tags": _detect_tags(raw_text + "\n" + corrected),
        "lines": [l.to_dict() for l in lines],
        "files": {
            "image": image_path.name,
            "image_raw": raw_image_path.name,
            "txt": txt_path.name,
            "docx": docx_path.name,
            "pdf": pdf_path.name,
        },
    }
    manifest_path = out_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    return ExportResult(
        lecture_id=lecture_id,
        captured_at=now.isoformat(),
        dir=str(out_dir),
        image_path=str(image_path),
        raw_image_path=str(raw_image_path),
        txt_path=str(txt_path),
        docx_path=str(docx_path),
        pdf_path=str(pdf_path),
        manifest_path=str(manifest_path),
    )
